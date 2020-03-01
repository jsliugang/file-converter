from gevent import monkey
monkey.patch_all()
from flask import Flask
from flask import request
from gevent import pywsgi
import os   
from docx import Document
from docx.oxml.ns import qn

app = Flask(__name__)

import logging
from logging import handlers
class Logger():
    def __init__(self):
        logging.basicConfig(level = logging.INFO,format = '%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        self.logger = logging.getLogger('fileconvert')
    def info(self,str):
        self.logger.info(str) 

def get_file_type(filename):
    '''
    获取文件的类型
    '''
    if '.pdf' == filename[-4:].lower():
        orginname=filename[:-4]
        return {'type':'pdf','orginname':orginname}
    elif '.doc' == filename[-4:].lower():
        orginname=filename[:-4]
        return {'type':'doc','orginname':orginname}
    elif '.docx' == filename[-5:].lower():
        orginname=filename[:-5]
        return {'type':'doc','orginname':orginname}
    
    elif '.ppt' == filename[-4:].lower():
        orginname=filename[:-4]
        return {'type':'ppt','orginname':orginname}
    elif '.pptx' == filename[-5:].lower():
        orginname=filename[:-5]
        return {'type':'ppt','orginname':orginname}
    
    elif '.xls' == filename[-4:].lower():
        orginname=filename[:-4]
        return {'type':'xls','orginname':orginname}
    elif '.xlsx' == filename[-5:].lower():
        orginname=filename[:-5]
        return {'type':'xls','orginname':orginname}
        
    elif '.txt' == filename[-4:].lower():
        orginname=filename[:-4]
        return {'type':'txt','orginname':orginname}
    
    else:
        return {'type':'None','orginname':'None'}
    
def txt2docx(infile,outfile):
    """
    转换txt文件为docx文件
    """
    try:
        with open(infile, 'r', encoding='gbk') as f:
            content = f.read()
    except:
        try:
            with open(infile, 'r', encoding='utf-8') as f:
                content = f.read()
        except:
            content = ''
        
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # run = document.add_heading('',level=3).add_run(title)
    # run.font.name=u'微软雅黑'
    # run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑') 
    
    # document.add_paragraph('')
    document.add_paragraph(content)

    document.save(outfile) 
 
   
@app.route('/convert',methods=["POST"])
def convert():
    oldname=request.values.get('file')
    temp=get_file_type(oldname)
    type,orgin=temp['type'],temp['orginname']
    
    # 转换doc为pdf
    if type=='doc' or type=='ppt' or type=='xls':
        os.system('unoconv -f pdf ./data/'+oldname)
        return 'Convert done.'
    elif type=='txt':
        # 由于centos中文编码不正常 暂时转换成docx处理 
        txt2docx('./data/'+oldname,'./data/'+orgin+'.docx')
        os.system('unoconv -f pdf ./data/'+orgin+'.docx')
        os.remove('./data/'+orgin+'.docx')
        return 'Convert done.'
    
    else:
        return 'Not support the file type.'
        
server = pywsgi.WSGIServer(('0.0.0.0', 9875), app)
server.serve_forever()